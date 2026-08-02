"""FormatSpec-aligned Word formatting for Free myself."""

from __future__ import annotations

import base64
import io
import re
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from docx.table import Table
from docx.text.paragraph import Paragraph

STYLE_CATALOG: dict[str, Any] = {
    "roles": [
        "title",
        "heading1",
        "heading2",
        "heading3",
        "heading4",
        "body",
        "bibliography",
    ],
    "fontEastAsia": ["黑体", "宋体", "楷体", "仿宋", "微软雅黑"],
    "fontAscii": ["Times New Roman", "Arial", "Calibri"],
    "align": ["left", "center", "right", "both"],
    "colors": ["#000000", "#333333", "#C00000"],
    "fontSizePt": {"min": 8, "max": 36, "step": 0.5},
    "lineSpacing": {"min": 1.0, "max": 3.0, "step": 0.25},
    "firstLineIndentChars": {"min": 0, "max": 4, "step": 1},
    "marginCm": {"min": 1.0, "max": 5.0, "step": 0.1},
}

DEFAULT_MATCH_RULES: list[dict[str, str]] = [
    {"role": "heading1", "pattern": r"^第.+[章节部]"},
    {"role": "heading2", "pattern": r"^[一二三四五六七八九十]+[、.．]"},
    {"role": "heading2", "pattern": r"^\d+[、.](?!\d)"},
    {"role": "heading4", "pattern": r"^\d+\.\d+\.\d+"},
    {"role": "heading3", "pattern": r"^\d+\.\d+(?!\.\d)"},
    {"role": "heading3", "pattern": r"^[（(][一二三四五六七八九十\d]+[）)]"},
    {"role": "bibliography", "pattern": r"^参考文献"},
]

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

ROLE_TO_WORD_STYLE = {
    "title": "Title",
    "heading1": "Heading 1",
    "heading2": "Heading 2",
    "heading3": "Heading 3",
    "heading4": "Heading 4",
    "body": "Normal",
    "quote": "Intense Quote",
    "bibliography": "Normal",
    "caption": "Caption",
    "code": "Normal",
    "footer": "Normal",
}

STYLE_NAME_TO_ROLE = {
    "title": "title",
    "heading 1": "heading1",
    "heading1": "heading1",
    "heading 2": "heading2",
    "heading2": "heading2",
    "heading 3": "heading3",
    "heading3": "heading3",
    "heading 4": "heading4",
    "heading4": "heading4",
    "quote": "quote",
    "intense quote": "quote",
    "caption": "caption",
}


def styles_by_role(spec: dict[str, Any]) -> dict[str, dict[str, Any]]:
    out: dict[str, dict[str, Any]] = {}
    for s in spec.get("styles") or []:
        role = s.get("role")
        if role:
            out[role] = s
    return out


def pattern_to_regex(pattern: str, flags: str = "") -> re.Pattern[str]:
    p = (pattern or "").strip()
    flag = re.IGNORECASE if "i" in (flags or "").lower() else 0
    if not p:
        return re.compile(r"^$", flag)
    if p.startswith("^") or "\\" in p or any(c in p for c in "[]()+?{}|"):
        return re.compile(p, flag)
    escaped = re.escape(p).replace(r"\*", ".*")
    return re.compile("^" + escaped, flag)


def guess_role(
    paragraph: Paragraph,
    text: str,
    content_index: int,
    match_rules: list[dict[str, Any]] | None = None,
) -> str:
    """Guess role. ``content_index`` is 0-based among non-empty paragraphs."""
    name = (paragraph.style.name if paragraph.style else "Normal") or "Normal"
    key = name.strip().lower()
    if key in STYLE_NAME_TO_ROLE:
        return STYLE_NAME_TO_ROLE[key]

    stripped = text.strip()
    if not stripped:
        return "body"

    rules = match_rules if match_rules is not None else DEFAULT_MATCH_RULES
    for rule in rules:
        role = rule.get("role")
        pattern = rule.get("pattern") or ""
        if not role or not pattern:
            continue
        try:
            rx = pattern_to_regex(pattern, rule.get("flags") or "")
            if rx.search(stripped):
                return str(role)
        except re.error:
            continue

    if content_index == 0 and len(stripped) < 80:
        return "title"
    if re.match(r"^第[一二三四五六七八九十百千0-9]+[章节部]", stripped):
        return "heading1"
    if re.match(r"^[一二三四五六七八九十]+[、.．]", stripped):
        return "heading2"
    if stripped.startswith("参考文献") or stripped.lower().startswith("reference"):
        return "bibliography"
    if stripped.startswith(">") or stripped.startswith("「"):
        return "quote"
    return "body"


def _parse_hex(color: str) -> tuple[int, int, int]:
    c = (color or "#000000").strip().lstrip("#")
    if len(c) != 6:
        c = "000000"
    return int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)


def _hex_to_rgb(color: str) -> RGBColor:
    r, g, b = _parse_hex(color)
    return RGBColor(r, g, b)


def _set_rfonts(run, font_ascii: str, font_east: str) -> None:
    run.font.name = font_ascii
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_ascii)
    rFonts.set(qn("w:hAnsi"), font_ascii)
    rFonts.set(qn("w:eastAsia"), font_east)
    rFonts.set(qn("w:cs"), font_ascii)


def _force_run_color(run, color: str) -> None:
    """Force RGB color and strip theme colors that make Heading blue."""
    r, g, b = _parse_hex(color)
    run.font.color.rgb = RGBColor(r, g, b)
    rPr = run._element.get_or_add_rPr()
    for child in list(rPr):
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local in ("color",) or "theme" in local.lower():
            rPr.remove(child)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{r:02X}{g:02X}{b:02X}")
    rPr.append(color_el)


def set_run_font(
    run,
    style: dict[str, Any],
    *,
    preserve_emphasis: bool = False,
) -> None:
    """Apply role fonts/size/color. Optionally keep per-run bold/italic."""
    font_ascii = style.get("fontAscii") or "Times New Roman"
    font_east = style.get("fontEastAsia") or "宋体"
    size = float(style.get("fontSizePt") or 12)
    if not preserve_emphasis:
        run.font.bold = bool(style.get("bold", False))
        run.font.italic = bool(style.get("italic", False))
    _set_rfonts(run, font_ascii, font_east)
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    for tag in ("w:sz", "w:szCs"):
        existing = rPr.find(qn(tag))
        if existing is not None:
            rPr.remove(existing)
        node = OxmlElement(tag)
        node.set(qn("w:val"), str(int(round(size * 2))))
        rPr.append(node)
    _force_run_color(run, style.get("color") or "#000000")


def _run_has_drawing(run) -> bool:
    el = run._element
    return bool(
        el.findall(".//" + qn("w:drawing"))
        or el.findall(".//" + qn("w:pict"))
        or el.findall(".//" + qn("w:object"))
    )


def apply_paragraph_style(
    paragraph: Paragraph,
    style: dict[str, Any],
    *,
    role: str,
    doc: Document | None = None,
) -> None:
    """Apply paragraph + run formatting without collapsing runs (keeps images/links)."""
    word_style_name = ROLE_TO_WORD_STYLE.get(role, "Normal")
    if doc is not None:
        try:
            paragraph.style = doc.styles[word_style_name]
        except KeyError:
            try:
                paragraph.style = doc.styles["Normal"]
            except KeyError:
                pass

    align = style.get("align") or "both"
    paragraph.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

    pf = paragraph.paragraph_format
    pf.space_before = Pt(float(style.get("spaceBeforePt") or 0))
    pf.space_after = Pt(float(style.get("spaceAfterPt") or 0))

    line = float(style.get("lineSpacing") or 1.5)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line

    size = float(style.get("fontSizePt") or 12)
    first_chars = float(style.get("firstLineIndentChars") or 0)
    hanging_chars = float(style.get("hangingIndentChars") or 0)
    char_twips = size * 20
    if hanging_chars > 0:
        pf.first_line_indent = Twips(int(-hanging_chars * char_twips))
        pf.left_indent = Twips(int(hanging_chars * char_twips))
    elif first_chars > 0:
        pf.first_line_indent = Twips(int(first_chars * char_twips))
        pf.left_indent = Twips(0)
    else:
        pf.first_line_indent = Twips(0)
        pf.left_indent = Twips(0)

    # Headings force bold from role; body keeps inline bold/italic
    preserve = role in ("body", "bibliography", "caption", "footer", "code")
    runs = list(paragraph.runs)
    if not runs:
        text = (paragraph.text or "").strip()
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, style, preserve_emphasis=False)
        return

    for run in runs:
        if _run_has_drawing(run) and not (run.text or "").strip():
            continue  # leave image-only runs alone
        set_run_font(run, style, preserve_emphasis=preserve)


def _patch_style_definition(doc: Document, word_name: str, style: dict[str, Any]) -> None:
    try:
        st = doc.styles[word_name]
    except KeyError:
        return
    if st.type != WD_STYLE_TYPE.PARAGRAPH:
        return
    font = st.font
    font_ascii = style.get("fontAscii") or "Times New Roman"
    font.size = Pt(float(style.get("fontSizePt") or 12))
    font.bold = bool(style.get("bold", False))
    font.italic = bool(style.get("italic", False))
    font.name = font_ascii
    try:
        font.color.rgb = _hex_to_rgb(style.get("color") or "#000000")
    except Exception:  # noqa: BLE001
        pass
    try:
        rPr = st.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), font_ascii)
        rFonts.set(qn("w:hAnsi"), font_ascii)
        rFonts.set(qn("w:eastAsia"), style.get("fontEastAsia") or "宋体")
        # strip theme color on style
        for child in list(rPr):
            if "theme" in child.tag.lower():
                rPr.remove(child)
        existing = rPr.find(qn("w:color"))
        if existing is not None:
            rPr.remove(existing)
        rgb = _parse_hex(style.get("color") or "#000000")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
        rPr.append(color_el)
    except Exception:  # noqa: BLE001
        pass
    align = style.get("align")
    if align and align in ALIGN_MAP:
        st.paragraph_format.alignment = ALIGN_MAP[align]
    line = float(style.get("lineSpacing") or 0)
    if line:
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        st.paragraph_format.line_spacing = line


def apply_document_style_defs(doc: Document, by_role: dict[str, dict[str, Any]]) -> None:
    for role, word_name in ROLE_TO_WORD_STYLE.items():
        if role in by_role:
            _patch_style_definition(doc, word_name, by_role[role])


def apply_margins(doc: Document, spec: dict[str, Any]) -> None:
    margin = (spec.get("meta") or {}).get("marginCm") or {}
    top = float(margin.get("top", 2.54))
    bottom = float(margin.get("bottom", 2.54))
    left = float(margin.get("left", 3.17))
    right = float(margin.get("right", 3.17))
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _set_cell_shading(cell, fill_hex: str) -> None:
    fill = (fill_hex or "").strip().lstrip("#")
    if not fill:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill.upper())
    tcPr.append(shd)


def _set_table_borders(table: Table, enabled: bool) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    val = "single" if enabled else "nil"
    sz = "4" if enabled else "0"
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def apply_tables(doc: Document, table_spec: dict[str, Any] | None, doc_ref: Document) -> int:
    if not table_spec:
        table_spec = {}
    count = 0
    header_style = {
        "fontEastAsia": table_spec.get("headerFontEastAsia") or "黑体",
        "fontAscii": "Times New Roman",
        "fontSizePt": float(table_spec.get("headerFontSizePt") or 12),
        "bold": bool(table_spec.get("headerBold", True)),
        "color": table_spec.get("headerColor") or "#000000",
        "align": table_spec.get("align") or "center",
        "lineSpacing": 1.15,
        "firstLineIndentChars": 0,
        "hangingIndentChars": 0,
        "spaceBeforePt": 0,
        "spaceAfterPt": 0,
    }
    body_style = {
        "fontEastAsia": table_spec.get("bodyFontEastAsia") or "宋体",
        "fontAscii": "Times New Roman",
        "fontSizePt": float(table_spec.get("bodyFontSizePt") or 10.5),
        "bold": False,
        "color": table_spec.get("bodyColor") or "#000000",
        "align": table_spec.get("align") or "center",
        "lineSpacing": 1.15,
        "firstLineIndentChars": 0,
        "hangingIndentChars": 0,
        "spaceBeforePt": 0,
        "spaceAfterPt": 0,
    }
    shading = table_spec.get("headerShading") or ""
    borders = bool(table_spec.get("borders", True))

    for table in doc.tables:
        _set_table_borders(table, borders)
        for r_idx, row in enumerate(table.rows):
            is_header = r_idx == 0
            style = header_style if is_header else body_style
            for cell in row.cells:
                if is_header and shading:
                    _set_cell_shading(cell, shading)
                for para in cell.paragraphs:
                    if not (para.text or "").strip():
                        # still apply alignment on empty?
                        continue
                    apply_paragraph_style(para, style, role="body", doc=doc_ref)
                    count += 1
    return count


def format_existing_docx(
    docx_bytes: bytes, spec: dict[str, Any]
) -> tuple[bytes, dict[str, Any]]:
    doc = Document(io.BytesIO(docx_bytes))
    apply_margins(doc, spec)
    by_role = styles_by_role(spec)
    body_style = by_role.get("body") or {
        "fontEastAsia": "宋体",
        "fontAscii": "Times New Roman",
        "fontSizePt": 12,
        "color": "#000000",
        "align": "both",
        "lineSpacing": 1.5,
        "firstLineIndentChars": 2,
    }
    if "color" not in body_style:
        body_style = {**body_style, "color": "#000000"}
    apply_document_style_defs(doc, by_role)

    match_rules = spec.get("matchRules")
    if not match_rules:
        match_rules = DEFAULT_MATCH_RULES

    role_counts: dict[str, int] = {}
    applied = 0
    content_index = 0
    for para in doc.paragraphs:
        text = para.text or ""
        if not text.strip():
            continue
        role = guess_role(para, text, content_index, match_rules)
        content_index += 1
        style = {**(by_role.get(role) or body_style)}
        if not style.get("color"):
            style["color"] = "#000000"
        apply_paragraph_style(para, style, role=role, doc=doc)
        role_counts[role] = role_counts.get(role, 0) + 1
        applied += 1

    table_cells = apply_tables(doc, spec.get("table"), doc)
    applied += table_cells

    out = io.BytesIO()
    doc.save(out)
    margin = (spec.get("meta") or {}).get("marginCm") or {}
    summary = {
        "paragraphsStyled": applied,
        "rolesUsed": role_counts,
        "tableCellsStyled": table_cells,
        "marginCm": margin,
        "bodyFont": body_style.get("fontEastAsia"),
        "bodySizePt": body_style.get("fontSizePt"),
        "forcedColor": body_style.get("color") or "#000000",
    }
    return out.getvalue(), summary


def _add_styled_paragraph(
    doc: Document, text: str, style: dict[str, Any], role: str
) -> None:
    para = doc.add_paragraph()
    word_style_name = ROLE_TO_WORD_STYLE.get(role, "Normal")
    try:
        para.style = doc.styles[word_style_name]
    except KeyError:
        pass

    style = {**style, "color": style.get("color") or "#000000"}
    align = style.get("align") or "both"
    para.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = para.paragraph_format
    pf.space_before = Pt(float(style.get("spaceBeforePt") or 0))
    pf.space_after = Pt(float(style.get("spaceAfterPt") or 0))
    line = float(style.get("lineSpacing") or 1.5)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    size = float(style.get("fontSizePt") or 12)
    first_chars = float(style.get("firstLineIndentChars") or 0)
    hanging_chars = float(style.get("hangingIndentChars") or 0)
    char_twips = size * 20
    if hanging_chars > 0:
        pf.first_line_indent = Twips(int(-hanging_chars * char_twips))
        pf.left_indent = Twips(int(hanging_chars * char_twips))
    elif first_chars > 0:
        pf.first_line_indent = Twips(int(first_chars * char_twips))
    else:
        pf.first_line_indent = Twips(0)

    run = para.add_run(text)
    set_run_font(run, style)


def markdown_to_docx(
    markdown: str, spec: dict[str, Any]
) -> tuple[bytes, dict[str, Any]]:
    by_role = styles_by_role(spec)
    mapping = spec.get("mdMapping") or {}
    body = by_role.get("body") or {
        "fontEastAsia": "宋体",
        "fontAscii": "Times New Roman",
        "fontSizePt": 12,
        "color": "#000000",
        "align": "both",
        "lineSpacing": 1.5,
        "firstLineIndentChars": 2,
    }

    def role_style(role: str) -> dict[str, Any]:
        s = dict(by_role.get(role) or body)
        s["color"] = s.get("color") or "#000000"
        return s

    doc = Document()
    apply_margins(doc, spec)
    apply_document_style_defs(doc, by_role)
    if doc.paragraphs and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_buf: list[str] = []
    role_counts: dict[str, int] = {}

    def emit(text: str, role: str) -> None:
        _add_styled_paragraph(doc, text, role_style(role), role)
        role_counts[role] = role_counts.get(role, 0) + 1

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                emit("\n".join(code_buf), mapping.get("code") or "code")
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("#### "):
            emit(line[5:].strip(), mapping.get("h4") or "heading4")
        elif line.startswith("### "):
            emit(line[4:].strip(), mapping.get("h3") or "heading3")
        elif line.startswith("## "):
            emit(line[3:].strip(), mapping.get("h2") or "heading2")
        elif line.startswith("# "):
            if len(doc.paragraphs) == 0 and "title" in by_role:
                emit(line[2:].strip(), "title")
            else:
                emit(line[2:].strip(), mapping.get("h1") or "heading1")
        elif line.startswith("> "):
            emit(line[2:].strip(), mapping.get("blockquote") or "quote")
        else:
            emit(line.strip(), mapping.get("paragraph") or "body")
        i += 1

    if in_code and code_buf:
        emit("\n".join(code_buf), mapping.get("code") or "code")

    out = io.BytesIO()
    doc.save(out)
    summary = {
        "paragraphsStyled": sum(role_counts.values()),
        "rolesUsed": role_counts,
        "bodyFont": body.get("fontEastAsia"),
        "bodySizePt": body.get("fontSizePt"),
        "forcedColor": "#000000",
    }
    return out.getvalue(), summary


def bytes_to_b64(data: bytes) -> str:
    return base64.b64encode(data).decode("ascii")
