"""Tests: black color, match rules, tables."""

from __future__ import annotations

import io

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from formatter import format_existing_docx, guess_role, markdown_to_docx, prioritize_match_rules

LOUD_SPEC = {
    "meta": {
        "name": "夸张测试规范",
        "scene": "test",
        "paper": "A4",
        "marginCm": {"top": 2.0, "bottom": 2.0, "left": 4.0, "right": 2.5},
    },
    "styles": [
        {
            "role": "title",
            "fontEastAsia": "黑体",
            "fontAscii": "Arial",
            "fontSizePt": 22,
            "bold": True,
            "color": "#000000",
            "align": "center",
            "lineSpacing": 1.5,
            "firstLineIndentChars": 0,
        },
        {
            "role": "heading1",
            "fontEastAsia": "黑体",
            "fontAscii": "Arial",
            "fontSizePt": 18,
            "bold": True,
            "color": "#000000",
            "align": "left",
            "lineSpacing": 1.5,
        },
        {
            "role": "body",
            "fontEastAsia": "楷体",
            "fontAscii": "Arial",
            "fontSizePt": 18,
            "bold": False,
            "color": "#000000",
            "align": "both",
            "lineSpacing": 2.0,
            "firstLineIndentChars": 2,
        },
    ],
    "matchRules": [
        {"role": "heading1", "pattern": "^第.+章"},
        {"role": "heading2", "pattern": "^[一二三四五六七八九十]+、"},
    ],
    "table": {
        "headerBold": True,
        "headerFontEastAsia": "黑体",
        "headerFontSizePt": 12,
        "headerColor": "#000000",
        "headerShading": "D9D9D9",
        "bodyFontEastAsia": "宋体",
        "bodyFontSizePt": 10.5,
        "bodyColor": "#000000",
        "borders": True,
        "align": "center",
    },
    "mdMapping": {
        "h1": "heading1",
        "h2": "heading2",
        "h3": "heading3",
        "h4": "heading4",
        "paragraph": "body",
        "blockquote": "quote",
        "code": "code",
    },
}


def _sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("人工智能导论课程报告")
    doc.add_paragraph("第一章 背景")
    doc.add_paragraph("本节讨论大模型与文档排版。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头A"
    table.cell(0, 1).text = "表头B"
    table.cell(1, 0).text = "内容1"
    table.cell(1, 1).text = "内容2"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _run_color_val(run) -> str | None:
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return None
    color = rPr.find(qn("w:color"))
    if color is None:
        return None
    return color.get(qn("w:val"))


def test_format_docx_black_and_roles():
    out, summary = format_existing_docx(_sample_docx_bytes(), LOUD_SPEC)
    assert summary["rolesUsed"].get("heading1", 0) >= 1
    assert summary["forcedColor"] == "#000000"
    assert summary["tableCellsStyled"] >= 2

    doc = Document(io.BytesIO(out))
    assert abs(doc.sections[0].left_margin.cm - 4.0) < 0.05
    paras = [p for p in doc.paragraphs if p.text.strip()]
    assert paras[0].runs[0].font.size == Pt(22)
    assert paras[0].runs[0].font.color.rgb == RGBColor(0, 0, 0)
    assert _run_color_val(paras[0].runs[0]) == "000000"

    # Heading style must not leave theme blue
    h1 = next(p for p in paras if "第一章" in p.text)
    assert h1.runs[0].font.color.rgb == RGBColor(0, 0, 0)


def test_match_rules_heading1():
    doc = Document()
    p = doc.add_paragraph("第一章 背景")
    role = guess_role(p, p.text, 1, LOUD_SPEC["matchRules"])
    assert role == "heading1"


def test_custom_match_rule_beats_preset():
    """Custom pattern listed after presets still wins once prioritized."""
    doc = Document()
    p = doc.add_paragraph("第一章 特殊处理")
    rules = [
        {"role": "heading1", "pattern": r"^第.+[章节部]"},
        {"role": "heading2", "pattern": r"^第一章"},  # custom override
    ]
    assert prioritize_match_rules(rules)[0]["role"] == "heading2"
    assert guess_role(p, p.text, 1, rules) == "heading2"


def test_md_to_docx_loud_body():
    md = "# 人工智能导论\n\n## 背景\n\n本节讨论排版。\n"
    out, summary = markdown_to_docx(md, LOUD_SPEC)
    assert summary["paragraphsStyled"] >= 2
    doc = Document(io.BytesIO(out))
    assert abs(doc.sections[0].left_margin.cm - 4.0) < 0.05
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "人工智能导论" in texts[0]
    assert doc.paragraphs[0].runs[0].font.color.rgb == RGBColor(0, 0, 0)


def test_title_skips_leading_blank_paragraphs():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("真正的标题")
    doc.add_paragraph("正文第一段")
    buf = io.BytesIO()
    doc.save(buf)

    out, summary = format_existing_docx(buf.getvalue(), LOUD_SPEC)
    assert summary["rolesUsed"].get("title", 0) == 1
    result = Document(io.BytesIO(out))
    title = next(p for p in result.paragraphs if p.text.strip() == "真正的标题")
    assert title.runs[0].font.size == Pt(22)


def test_format_preserves_inline_runs():
    """Body paragraphs must keep multiple runs and inline bold."""
    doc = Document()
    doc.add_paragraph("报告标题")  # first para → title role
    p = doc.add_paragraph()
    p.add_run("普通文字")
    bold = p.add_run("加粗片段")
    bold.bold = True
    p.add_run("继续普通")
    buf = io.BytesIO()
    doc.save(buf)

    out, _ = format_existing_docx(buf.getvalue(), LOUD_SPEC)
    result = Document(io.BytesIO(out))
    body = next(para for para in result.paragraphs if "普通文字" in para.text)
    assert len(body.runs) >= 3
    assert body.text == "普通文字加粗片段继续普通"
    bold_run = next(r for r in body.runs if r.text == "加粗片段")
    assert bold_run.bold is True
    assert bold_run.font.size == Pt(18)
    assert _run_color_val(bold_run) == "000000"
